import pdfplumber
import pytesseract
from pdf2image import convert_from_path
import re
from difflib import SequenceMatcher
from dataclasses import dataclass
from pathlib import Path

# ---------------------------------------------------------
# Configure paths (Windows)
# ---------------------------------------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\Program Files\poppler-25.11.0\Library\bin"  # ← ADD THIS (adjust to your install path)


# ---------------------------------------------------------
# ParsedDocument — structured output with confidence
# ---------------------------------------------------------
@dataclass
class ParsedDocument:
    text: str
    extraction_method: str   # "pdfplumber" | "ocr"
    confidence: float        # 0.0 – 1.0 (used by ranker for flagging)
    page_count: int


# =========================================================
# 1) COLUMN-AWARE EXTRACTION (pdfplumber)
# =========================================================
def extract_text_pdfplumber(pdf_path):
    """
    Column-aware extraction using pdfplumber.
    Splits page into left/right halves and reads each
    column independently to preserve reading order.
    """
    full_text = ""

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:

                page_width = page.width

                # ── Try column-aware split first ──────────────────
                left_bbox  = (0,            0, page_width / 2, page.height)
                right_bbox = (page_width / 2, 0, page_width,   page.height)

                left_crop  = page.crop(left_bbox)
                right_crop = page.crop(right_bbox)

                left_text  = left_crop.extract_text(x_tolerance=3, y_tolerance=3) or ""
                right_text = right_crop.extract_text(x_tolerance=3, y_tolerance=3) or ""

                # ── Decide: single-column or two-column ──────────
                # If right column has very little text, treat as single column
                if len(right_text.strip()) < 50:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                else:
                    # True two-column: read left then right
                    page_text = left_text + "\n" + right_text

                full_text += page_text + "\n"

    except Exception as e:
        print(f"PDFPlumber Error: {e}")

    return full_text


# =========================================================
# 2) OCR FALLBACK
# =========================================================
def extract_text_ocr(pdf_path):
    """
    OCR using Tesseract with poppler path explicitly set.
    """
    text = ""

    try:
        images = convert_from_path(
            pdf_path,
            dpi=300,
            poppler_path=POPPLER_PATH   # ← explicit path prevents silent failures
        )

        for img in images:
            # PSM 6 = assume a single uniform block of text
            # Helps with resume layouts
            custom_config = r'--oem 3 --psm 6'
            ocr_text = pytesseract.image_to_string(img, lang="eng", config=custom_config)
            text += ocr_text + "\n"

    except Exception as e:
        print(f"OCR Error: {e}")

    return text


# =========================================================
# 3) OCR NOISE CLEANER (NEW)
# =========================================================
OCR_CORRECTIONS = [
    (r"[''`](\d{2})\b",                    r"20\1"),       # Aug'22 → Aug 2022
    (r"\b(\d{4})\s*[-–]\s*[Pp]re[sc]ent", r"\1 - Present"),
    (r"\bCun?ent\b",                        "Current"),     # OCR misread
    (r"\bPiesent\b",                        "Present"),     # OCR misread
    (r"\bTill [Dd]ate\b",                   "Present"),
    (r"(\d)\s+(\d{3})\b",                   r"\1\2"),       # 2 023 → 2023
    (r"\brn\b",                             "m"),           # rn → m (font misread)
    (r"\bI{2,}\b",                          ""),            # III artifacts
]

def clean_ocr_noise(text: str) -> str:
    for pattern, replacement in OCR_CORRECTIONS:
        text = re.sub(pattern, replacement, text)
    return text


# =========================================================
# 4) BROKEN LAYOUT DETECTION + CONFIDENCE SCORE
# =========================================================
def assess_text_quality(text: str) -> tuple[bool, float]:
    """
    Returns (is_broken: bool, confidence: float 0-1)
    Confidence is passed through the pipeline for flagging.
    """
    if not text or len(text.strip()) == 0:
        return True, 0.0

    words = text.split()
    if len(words) < 30:
        return True, 0.1

    # Single character word ratio
    single_chars   = sum(1 for w in words if len(w) == 1)
    ratio_single   = single_chars / len(words)

    # Symbol noise ratio
    symbol_count   = len(re.findall(r"[◆•|■►▪]", text))
    symbol_ratio   = symbol_count / max(len(words), 1)

    # Short line ratio
    lines          = text.splitlines()
    short_lines    = sum(1 for l in lines if 0 < len(l.strip()) <= 4)
    line_ratio     = short_lines / max(len(lines), 1)

    # Alphanumeric density (garbled OCR has low density)
    alnum_ratio    = sum(c.isalnum() or c.isspace() for c in text) / len(text)

    # ── Compute confidence ────────────────────────────────
    penalty = (
        ratio_single  * 0.4 +
        symbol_ratio  * 0.3 +
        line_ratio    * 0.2 +
        (1 - alnum_ratio) * 0.1
    )
    confidence = round(max(0.0, min(1.0, 1.0 - penalty)), 2)

    is_broken = ratio_single > 0.18 or symbol_ratio > 0.02 or line_ratio > 0.30

    return is_broken, confidence


# =========================================================
# 5) TEXT CLEANER
# =========================================================
def clean_text(text: str) -> str:
    """
    Clean and deduplicate resume text.
    Fixed: short-line threshold raised to avoid breaking
    skill lists, company names, date ranges.
    """
    # Remove cid encoded artifacts
    text = re.sub(r'\(cid:\d+\)', '•', text)

    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\r', '\n', text)

    raw_lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 0]

    paragraphs = []
    buffer = ""

    for line in raw_lines:
        # ── FIX: raised threshold 35 → 60 ────────────────
        # 35 was too aggressive — broke skill lines, date ranges
        # 60 still catches true section headers (EDUCATION, SKILLS etc.)
        if len(line) < 60 and line.isupper():
            # Only break on ALL-CAPS short lines (section headers)
            if buffer:
                paragraphs.append(buffer.strip())
                buffer = ""
            paragraphs.append(line)
            continue

        buffer += " " + line

    if buffer:
        paragraphs.append(buffer.strip())

    # ── Deduplicate paragraphs ────────────────────────────
    unique_paragraphs = []
    for para in paragraphs:
        duplicate = False
        for existing in unique_paragraphs:
            similarity = SequenceMatcher(None, para.lower(), existing.lower()).ratio()
            if similarity > 0.82:
                duplicate = True
                break
        if not duplicate:
            unique_paragraphs.append(para)

    return "\n\n".join(unique_paragraphs)


# =========================================================
# 6) MAIN FUNCTION — returns ParsedDocument now
# =========================================================
def extract_text_from_pdf(pdf_path) -> ParsedDocument:
    """
    Hybrid Resume Parsing Pipeline.
    Now returns ParsedDocument with confidence score
    instead of raw string — used by ranker for flagging.
    """
    import os
    page_count = 0

    try:
        with pdfplumber.open(pdf_path) as pdf:
            page_count = len(pdf.pages)
    except Exception:
        pass

    # Step 1 — Column-aware extraction
    text = extract_text_pdfplumber(pdf_path)

    # Step 2 — Assess quality
    is_broken, confidence = assess_text_quality(text)
    method = "pdfplumber"

    # Step 3 — OCR fallback if broken
    if is_broken:
        print(f"⚠ Broken layout detected → Switching to OCR... [{os.path.basename(pdf_path)}]")
        text = extract_text_ocr(pdf_path)
        text = clean_ocr_noise(text)          # ← NEW: clean OCR noise before anything else
        method = "ocr"
        _, confidence = assess_text_quality(text)   # re-assess after OCR

    # Step 4 — Clean and deduplicate
    text = clean_text(text)

    return ParsedDocument(
        text=text,
        extraction_method=method,
        confidence=confidence,
        page_count=page_count
    )


# =========================================================
# 7) WORD DOCUMENT PARSER (.docx)
# =========================================================
def extract_text_from_docx(docx_path: str) -> ParsedDocument:
    """
    Extracts text from a .docx Word document.
    Covers paragraphs, tables, and section headers/footers.
    Returns ParsedDocument with extraction_method="docx".
    Note: python-docx has no page count API; page_count is set to 1.
    """
    from docx import Document

    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                if not hf.is_linked_to_previous:
                    for para in hf.paragraphs:
                        t = para.text.strip()
                        if t:
                            parts.append(t)
            except Exception:
                pass

    raw_text = "\n".join(parts)
    text = clean_text(raw_text)
    _, confidence = assess_text_quality(text)

    return ParsedDocument(
        text=text,
        extraction_method="docx",
        confidence=confidence,
        page_count=1,
    )


# =========================================================
# 8) IMAGE PARSER (.jpg, .png, .bmp, .tiff, etc.)
# =========================================================
def extract_text_from_image(image_path: str) -> ParsedDocument:
    """
    Extracts text from a resume image via Tesseract OCR.
    Preprocessing: grayscale conversion + Otsu binarization.
    Returns ParsedDocument with extraction_method="image_ocr".
    """
    from PIL import Image, ImageOps, ImageFilter

    img = Image.open(image_path)

    if img.mode != "L":
        img = img.convert("L")

    # Upscale small images to give Tesseract better resolution
    w, h = img.size
    if w < 1200:
        scale = 1200 / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # Light sharpening improves character edges
    img = img.filter(ImageFilter.SHARPEN)

    custom_config = r"--oem 3 --psm 6"
    raw_text = pytesseract.image_to_string(img, lang="eng", config=custom_config)

    text = clean_ocr_noise(raw_text)
    text = clean_text(text)
    _, confidence = assess_text_quality(text)

    return ParsedDocument(
        text=text,
        extraction_method="image_ocr",
        confidence=confidence,
        page_count=1,
    )


# =========================================================
# 9) UNIVERSAL ROUTER — dispatches by file extension
# =========================================================
_EXTENSION_MAP = {
    ".pdf":  extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".jpg":  extract_text_from_image,
    ".jpeg": extract_text_from_image,
    ".png":  extract_text_from_image,
    ".bmp":  extract_text_from_image,
    ".tiff": extract_text_from_image,
    ".tif":  extract_text_from_image,
    ".webp": extract_text_from_image,
}

SUPPORTED_EXTENSIONS = set(_EXTENSION_MAP.keys())


def extract_text(file_path: str) -> ParsedDocument:
    """
    Universal entry point. Routes to the correct parser based on extension.
    Raises ValueError for unsupported file types.
    """
    ext = Path(file_path).suffix.lower()
    parser = _EXTENSION_MAP.get(ext)
    if parser is None:
        raise ValueError(
            f"Unsupported file format '{ext}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return parser(file_path)